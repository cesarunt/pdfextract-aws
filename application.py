import os, re
from flask import Flask, render_template, request, redirect, make_response, jsonify, send_file
from utils.config import cfg
from utils.handle_files import allowed_file, allowed_file_filesize, get_viewProcess_CPU
from werkzeug.utils import secure_filename
from scripts.split import pdf_remove, pdf_splitter
from scripts.process import pdf_process
from datetime import datetime
from docx import Document
from docx.shared import Pt 
from fold_to_ascii import fold

application = Flask(__name__)

application.jinja_env.auto_reload = True
application.config["TEMPLATES_AUTO_RELOAD"] = True
application.config['MAX_CONTENT_LENGTH'] = cfg.FILES.MAX_CONTENT_LENGTH 
application.config['UPLOAD_EXTENSIONS']  = cfg.FILES.UPLOAD_EXTENSIONS
application.config['SINGLE_UPLOAD']      = cfg.FILES.SINGLE_UPLOAD
application.config['SINGLE_SPLIT']       = cfg.FILES.SINGLE_SPLIT
application.config['SINGLE_OUTPUT']      = cfg.FILES.SINGLE_OUTPUT
application.config['SINGLE_FORWEB']      = cfg.FILES.SINGLE_FORWEB
application.config['MULTIPLE_UPLOAD']    = cfg.FILES.MULTIPLE_UPLOAD
application.config['MULTIPLE_SPLIT']     = cfg.FILES.MULTIPLE_SPLIT
application.config['MULTIPLE_OUTPUT']    = cfg.FILES.MULTIPLE_OUTPUT
application.config['MULTIPLE_FORWEB']    = cfg.FILES.MULTIPLE_FORWEB

# Allowed extension you can set your own
ALLOWED_EXTENSIONS = set(['PDF', 'pdf'])
ILLEGAL_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]")

def strip_illegal_xml_characters(s, default, base=10):
    # Compare the "invalid XML character range" numerically
    n = int(s, base)
    if n in (0xb, 0xc, 0xFFFE, 0xFFFF) or 0x0 <= n <= 0x8 or 0xe <= n <= 0x1F or 0xD800 <= n <= 0xDFFF:
        return ""
    return default

def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None

def validate_path(path):
    new_path = path
    path_split = path.split("/")
    filename = path_split[-1]

    if filename[0] == "_":
        i = 0
        for c in filename:
            if c != filename[0]:
                pos = i
                break
            i += 1
        new_path = "/".join(path_split[:-1])+"/"+filename[pos:]

    return new_path

def build_document_(title, text_pdf, language):
    # document = Document() 
    document.add_heading(title)

    keywords = "sample"
    texts = []

    # add a paragraphs
    if language != '' :
        for key, value in text_pdf:
            # doc = docx.Document()
            p = document.add_paragraph()

            patt = re.search(rf"\b{keywords}\b", value, re.IGNORECASE)
            # print("values")
            if patt != None:
                text_1 = value[:patt.start(0)]
                text_2 = keywords
                text_3 = value[patt.end(0):]
                texts = [tuple(["N", text_1]), tuple(["I", text_2]), tuple(["N", text_3])]
                band = True
            else:
                texts = [tuple([key, value])]

            for key_text, value_text in texts :
                # line = p.add_run(str(value.encode('utf-8').decode("utf-8")))
                try:
                    # p = document.add_paragraph(str(value.encode('utf-8').decode("utf-8")))
                    line = p.add_run(str(value_text.encode('utf-8').decode("utf-8")))
                except:
                    delete_paragraph(p)
                    p = document.add_paragraph()
                    html = value_text.encode("ascii", "xmlcharrefreplace").decode("utf-8")
                    html = re.sub(r"&#(\d+);?", lambda c: strip_illegal_xml_characters(c.group(1), c.group(0)), html)
                    html = re.sub(r"&#[xX]([0-9a-fA-F]+);?", lambda c: strip_illegal_xml_characters(c.group(1), c.group(0), base=16), html)
                    html = ILLEGAL_XML_CHARS_RE.sub("", html)
                    # p = document.add_paragraph(str(html.encode('utf-8').decode("utf-8")))
                    line = p.add_run(str(html.encode('utf-8').decode("utf-8")))

                if key_text == "B": line.bold = True
                if key_text == "I": line.bold = True; line.italic = True; line.font.size = Pt(12) #line.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
            
            texts = []
            # if key == "B":  p.add_run('bold').bold = True

    return document

def build_document(title, text_pdf, language):
    # document = Document() 
    document.add_heading(title)

    # add a paragraphs
    if language != '' :
        for key, value in text_pdf:
            # doc = docx.Document()
            p = document.add_paragraph()
            # line = p.add_run(str(value.encode('utf-8').decode("utf-8")))
            try:
                # p = document.add_paragraph(str(value.encode('utf-8').decode("utf-8")))
                line = p.add_run(str(value.encode('utf-8').decode("utf-8")))
            except:
                delete_paragraph(p)
                p = document.add_paragraph()
                html = value.encode("ascii", "xmlcharrefreplace").decode("utf-8")
                html = re.sub(r"&#(\d+);?", lambda c: strip_illegal_xml_characters(c.group(1), c.group(0)), html)
                html = re.sub(r"&#[xX]([0-9a-fA-F]+);?", lambda c: strip_illegal_xml_characters(c.group(1), c.group(0), base=16), html)
                html = ILLEGAL_XML_CHARS_RE.sub("", html)
                # p = document.add_paragraph(str(html.encode('utf-8').decode("utf-8")))
                # line = p.add_run(str(html.encode('utf-8').decode("utf-8")))
                line = p.add_run(str(html))

            if key == "B": line.bold = True
            
            # texts = []
            # if key == "B":  p.add_run('bold').bold = True

    return document

def allowed_files(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Index
@application.route('/')
def home():
    return render_template('index.html')

@application.route('/extract_one')
def extract_one():
    return render_template('extract_one.html')

@application.route('/extract_mul')
def extract_mul():
    return render_template('extract_mul.html')

@application.route('/report')
def report():
    return render_template('report.html')

# ----------------------------------- PDF EXTRACT ONE -----------------------------------
@application.route('/extract_one', methods=['POST'])
def extract_one_load():
    global file_pdf
    active_show = "active show"
    # _analytic = request.form.get('analytic')
    # 
    if request.method == "POST":
        # Code for One pdf
        if "filesize" in request.cookies:
            if not allowed_file_filesize(request.cookies["filesize"], application.config["MAX_CONTENT_LENGTH"]):
                # print("Filesize exceeded maximum limit")
                return redirect(request.url)
            file = request.files["file"]
            filesize = request.cookies.get("filesize")

            if file.filename == "":
                # print("No filename")
                return redirect(request.url)
            if int(filesize) > 0 :
                res = make_response(jsonify({"message": f"El PDF fue cargado con éxito."}), 200)
                # print("File uploaded")
                upload = True
            if allowed_file(file.filename, application.config["UPLOAD_EXTENSIONS"]):
                filename = secure_filename(file.filename)
                # file.save(os.path.join(application.config["UPLOAD_PATH_UP"], filename))
                file.save(os.path.join(application.config["SINGLE_UPLOAD"], filename))
                file_pdf = filename
                print("File saved")
                if (upload == True):
                    return res
            else:
                print("That file extension is not allowed")
                return redirect(request.url)

@application.route("/action_extract_one", methods=["GET", "POST"])
def action_extract_one():
    result_split = False
    text_pdf = []
    is_article = True

    if request.method == "POST":
        global file_pdf, document
        resultCPU = False
        result_save = None
        result_file_text = ""
        result_file_down = ""

        # Verify if posible to process
        if get_viewProcess_CPU() is True :
            
            document = Document() 

            file_pdf = fold(file_pdf)  
            path = os.path.join(application.config['SINGLE_UPLOAD'],file_pdf)
            fname = os.listdir(application.config['SINGLE_SPLIT']) #fname: List contain pdf documents names in folder
            # 1. SPLIT PDF
            # print("\n------------------- START SPLIT PROCESS -------------------")
            pdf_remove(fname, application.config['SINGLE_SPLIT'])       # Call pdf remove function
            result_split = pdf_splitter(path, application.config['SINGLE_SPLIT'])      # Call pdf splitter function
            
            if result_split == 0:
                result_save = False
                result_file_text = "No se completó el procesamiento."
                
            if result_split == 2:
                result_save = False
                result_file_text = "El PDF debe tener máximo " + str(cfg.FILES.MAX_NUMPAGES) + " páginas."
            
            if result_split == 1:
                # 2. Process PDF
                # print("\n------------------ START EXTRACT PROCESS ------------------")
                is_article, text_pdf, language = pdf_process(application.config['SINGLE_SPLIT'], application.config['SINGLE_OUTPUT'])           # Call pdf process function
                # print("Len TEXT PDF")
                document = build_document(file_pdf, text_pdf, language)
            # if result_split==True :
                if is_article == False:
                    result_save = 0
                    result_file_text = "Error cargando formato de Tesis ... \nMuy pronto estará disponible"
                elif len(text_pdf) > 1 :
                    file_save = application.config['SINGLE_OUTPUT']+'/background_'+file_pdf+'.docx'
                    document.save(file_save)
                    result_save = 1
                    result_file_text = file_pdf.split(".pdf")[0]
                    result_file_down = application.config['SINGLE_FORWEB']+'/background_'+file_pdf+'.docx'
                else:
                    result_save = 0
                    result_file_text = "Error en la carga del PDF"
            # else:
            #     result_save = 0
            #     result_file_text = "Error en el PDF"
        else:
            resultCPU = True
            result_file_text = "El servidor está procesando, debe esperar un momento."
    
    # print("File Download: " + str(result_file_down))
    # print(result_save)
    return render_template('extract_one.html', result_save=result_save, result_file_text=result_file_text, result_file_down=result_file_down)

@application.route("/close_extract_one/<source>")
def close_extract_one(source):
    url = "/" + source
    return redirect(url)

@application.route("/save_extract_one", methods=["POST"])
def save_extract_one():
    if request.method == "POST":
        down_image = request.form.get('down_image')
        
    return send_file(down_image, as_attachment=True)


# ----------------------------------- PDF EXTRACT MULTIPLE -----------------------------------
@application.route('/extract_mul', methods=['POST'])
def extract_mul_load():
    global file_pdfs
    upload = False
    file_pdfs = []
    result_split = []
    # _analytic = request.form.get('analytic')

    if request.method == "POST":
        # Code for multiple pdfs
        if 'files[]' not in request.files:
            print('No file part')
            return redirect(request.url)

        files = request.files.getlist('files[]')

        for file in files:
            file_pdfs.applicationend(file.filename)
            if file and allowed_file(file.filename, application.config["UPLOAD_EXTENSIONS"]):
                filename = secure_filename(file.filename)
                file.save(os.path.join(application.config['MULTIPLE_UPLOAD'], filename))
                upload = True
        
        if (upload == True):
            print('File(s) successfully uploaded')
            return render_template('extract_mul.html', resultLoad=upload)

@application.route("/action_extract_mul", methods=["GET", "POST"])
def action_extract_mul():
    text_pdf = []

    if request.method == "POST":
        global file_pdf, document
        resultCPU = False
        result_save = None
        result_file_text = "None"
        result_invalid_text = ""
        result_file_down = "None"
        result_valid = 0
        result_invalid = 0
        result_invalid_process = []
        # result_invalid_numpages = []

        # Verify if posible to process
        if get_viewProcess_CPU() is True :

            document = Document() 
            
            for filename in file_pdfs :
                filename = fold(filename)                
                path = os.path.join(application.config['MULTIPLE_UPLOAD'],filename)
                path = validate_path(path)
                path = path.replace('(','').replace(')','').replace(',','').replace('<','').replace('>','').replace('?','').replace('!','').replace('@','').replace('%','').replace('$','').replace('#','').replace('*','').replace('&','').replace(';','').replace('{','').replace('}','').replace('[','').replace(']','').replace('|','').replace('=','').replace('+','').replace(' ','_')
                fname = os.listdir(application.config['MULTIPLE_SPLIT'])

                # 1. SPLIT PDF
                # print("\n------------------- START SPLIT PROCESS -------------------")
                pdf_remove(fname, application.config['MULTIPLE_SPLIT'])       # Call pdf remove function

                result_split = pdf_splitter(path, application.config['MULTIPLE_SPLIT'])      # Call pdf splitter function

                if result_split == 0:
                    # result_save = False
                    result_invalid += 1
                    result_invalid_process.applicationend(filename + " ...NO se procesó")
                    # result_file_text = "No se logró procesar"
                if result_split == 2:
                    # result_save = False
                    result_invalid += 1
                    result_invalid_process.applicationend(filename + " ...supera el Nro páginas")
                    # result_file_text = "El PDF debe tener máximo " + str(cfg.FILES.MAX_NUMPAGES) + " páginas."
                if result_split == 1:
                    # 2. Process PDF
                    # print("\n------------------ START EXTRACT PROCESS ------------------")
                    _, text_pdf, language = pdf_process(application.config['MULTIPLE_SPLIT'], application.config['MULTIPLE_OUTPUT'])  # Call pdf process function
                    # print("Out web: " + application.config['MULTIPLE_FORWEB'])

                    if len(text_pdf) > 1 :
                        now = datetime.now()
                        document = build_document(filename, text_pdf, language)
                        file_save = application.config['MULTIPLE_OUTPUT']+'/background_multiple_'+now.strftime("%d%m%Y_%H%M%S")+'.docx'
                        document.save(file_save)
                        # result_save = True
                        result_valid += 1
                        result_file_text = "Antecedente Múltiple"
                        result_file_down = application.config['MULTIPLE_FORWEB']+'/background_multiple_'+now.strftime("%d%m%Y_%H%M%S")+'.docx'
                
                if result_valid > 0 :
                    result_save = True
                
                if result_invalid > 0 and result_valid == 0 :
                    result_save = False
                    result_file_text = "No fue posible procesar"
                
                if len(result_invalid_process) > 0 :
                    # result_save = False
                    result_invalid_text = (',  \n'.join(result_invalid_process))
                    
        else:
            result_file_text = "El servidor está procesando, espere un momento."
    
    return render_template('extract_mul.html', result_save=result_save, result_file_text=result_file_text, result_invalid_text=result_invalid_text, result_file_down=result_file_down)


@application.route("/close_extract_mul/<source>")
def close_extract_mul(source):
    url = "/" + source
    return redirect(url)

@application.route("/save_extract_mul", methods=["POST"])
def save_extract_mul():
    if request.method == "POST":
        down_image = request.form.get('down_image')
        
    return send_file(down_image, as_attachment=True)


# INIT PROJECT
# if __name__ == '__main__':
#     # start the flask application
#     application.run(debug=True, use_reloader=True)
#     # application.run(host="0.0.0.0", port="5000", debug=True, threaded=True, use_reloader=True)